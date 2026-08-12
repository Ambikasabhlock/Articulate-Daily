from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT=Path(__file__).resolve().parents[2]
OUT=ROOT / "docs" / "Articulate-Daily-Customer-Guide-v3.03.docx"
doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(.85); sec.bottom_margin=Inches(.8); sec.left_margin=Inches(.85); sec.right_margin=Inches(.85); sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
navy=RGBColor(21,49,45); teal=RGBColor(15,118,110); muted=RGBColor(98,118,113)
normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.font.color.rgb=navy; normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,before,after,color in [("Title",30,0,12,navy),("Heading 1",16,18,10,teal),("Heading 2",13,14,7,teal),("Heading 3",12,10,5,navy)]:
    style=doc.styles[name]; style.font.name="Calibri"; style.font.size=Pt(size); style.font.bold=True; style.font.color.rgb=color; style.paragraph_format.space_before=Pt(before); style.paragraph_format.space_after=Pt(after); style.paragraph_format.keep_with_next=True
for name in ["List Bullet","List Number"]:
    style=doc.styles[name]; style.font.name="Calibri"; style.font.size=Pt(11); style.paragraph_format.left_indent=Inches(.375); style.paragraph_format.first_line_indent=Inches(-.188); style.paragraph_format.space_after=Pt(4); style.paragraph_format.line_spacing=1.25

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); tcPr.append(shd)
def margins(cell):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for name,value in (("top",80),("start",120),("bottom",80),("end",120)):
        node=OxmlElement(f"w:{name}"); node.set(qn("w:w"),str(value)); node.set(qn("w:type"),"dxa"); tcMar.append(node)
def table(headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,(text,width) in enumerate(zip(headers,widths)):
        c=t.rows[0].cells[i]; c.width=Inches(width); c.text=text; shade(c,"0F766E"); margins(c)
        for run in c.paragraphs[0].runs: run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.size=Pt(10)
    trPr=t.rows[0]._tr.get_or_add_trPr(); repeat=OxmlElement("w:tblHeader"); repeat.set(qn("w:val"),"true"); trPr.append(repeat)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,(text,width) in enumerate(zip(row,widths)):
            cells[i].width=Inches(width); cells[i].text=text; margins(cells[i])
            if ri%2: shade(cells[i],"F6FAF9")
            for run in cells[i].paragraphs[0].runs: run.font.size=Pt(9.5); run.font.color.rgb=navy
def bullet(text): doc.add_paragraph(text,style="List Bullet")
def numbered(text): doc.add_paragraph(text,style="List Number")
def callout(title,text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.12); p.paragraph_format.right_indent=Inches(.12); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(8)
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),"EAF3F1"); pPr.append(shd); borders=OxmlElement("w:pBdr"); left=OxmlElement("w:left"); left.set(qn("w:val"),"single"); left.set(qn("w:sz"),"24"); left.set(qn("w:color"),"0F766E"); left.set(qn("w:space"),"8"); borders.append(left); pPr.append(borders)
    r=p.add_run(title+"  "); r.bold=True; r.font.color.rgb=teal; p.add_run(text)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("ARTICULATE DAILY"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=teal
p=doc.add_paragraph(style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Build clearer English every day")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Customer guide · Version 3.03"); r.font.size=Pt(14); r.font.color.rgb=muted
p=doc.add_paragraph("Articulate Daily helps you collect useful vocabulary, understand it in context, practise it repeatedly and use it with greater confidence at work and in daily life."); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
callout("The daily promise", "A focused routine that turns useful words into confident communication—without requiring an account.")
doc.add_page_break()

doc.add_heading("Start with a simple daily routine",1)
for item in ["Open Today and listen to the word of the day.","Save words connected to your real work, study or conversations.","Complete the short review, sentence and pronunciation targets.","Use one saved word in a real message, meeting or conversation.","Return tomorrow; repeated practice is more valuable than occasional long sessions."]: numbered(item)
doc.add_heading("Where your learning lives",2)
doc.add_paragraph("The app saves progress on your device first. Every core learning feature works without an account. Optional cloud sync can make the same profiles available on another device.")
table(["Mode","What it gives you","Remember"],[["Local-first","Immediate saving on this browser","Download backups; clearing browser data can remove the local copy"],["Cloud sync","Optional sign-in and cross-device state","Requires a configured Supabase project and internet"],["Offline","Vocabulary, practice and library","Dictionary and browser speech recognition may need internet"]],[1.2,2.7,2.6])

doc.add_heading("Explore every feature",1)
features=[["Today","Daily word, progress, review queue, weekly activity, streak and targets"],["Add a word","Meaning, example, category, tags, source, notes and difficulty"],["Extract words","On-device PDF, DOCX, image, text, Markdown, CSV or JSON processing"],["Practice","Sentences, review ratings, collocations, quizzes and UK English speech"],["Speaking Coach","Rule-based practice for meetings, interviews, data and disagreement"],["Library","Search, filter, favourite, pronounce, review, master or delete"],["Profiles & backup","Separate profiles, targets, backup, CSV, reports and optional sync"]]
table(["Area","How it helps"],features,[1.55,4.95])

doc.add_heading("Practise for memory, not collection",1)
doc.add_paragraph("A large word list is not the goal. Recall and use a smaller number of relevant words in meaningful contexts.")
for item in ["Write examples about your own projects and conversations.","Say a complete sentence aloud instead of repeating an isolated word.","Choose Again, Hard, Good or Easy honestly.","Record where you heard a word and when you want to use it.","Mark a word mastered only when you can recall and use it without help."]: bullet(item)
callout("Pronunciation note", "The recognition percentage compares the target with the browser transcript. It is useful feedback, not a clinical or phonetic assessment.")

doc.add_heading("Use documents safely and effectively",1)
for item in ["Choose material connected to something you are learning.","Wait while the app reads the file on your device.","Review suggestions and deselect irrelevant words.","Add your picks, then improve their meanings and examples."]: numbered(item)
doc.add_paragraph("Images are limited to 10 MB, other files to 20 MB, PDFs to 40 pages, and backup/CSV imports to 5 MB.")

doc.add_heading("Profiles, backup and recovery",1)
for item in ["Use profiles to separate learners or learning goals.","Download a JSON backup regularly and store it outside the browser.","Import only files you trust; backup import replaces current app state after confirmation.","Use CSV for vocabulary exchange and JSON for complete recovery.","Download a final backup before deleting data."]: bullet(item)
callout("Deletion protection", "Profile deletion requires its exact name. Deleting all app data requires two confirmations and DELETE ALL. When signed in, synced state is also removed.")

doc.add_heading("Privacy and responsible use",1)
doc.add_paragraph("Data stays on your device by default. Optional sync stores the same state in your authenticated Supabase account. Dictionary lookup sends the requested word to a public service. Document extraction uses bundled browser libraries. Browser speech recognition may use an online browser-vendor service.")
for item in ["Do not put passwords or authentication codes in notes.","Avoid confidential company, medical, financial or personal records.","Use HTTPS before enabling microphone features in production.","Sign out on shared devices."]: bullet(item)

doc.add_heading("Troubleshooting",1)
table(["Situation","What to do"],[["Voice recognition unavailable","Use current Chrome or Edge and confirm microphone permission"],["Speech does not sound British","Check installed system voices; an available en-GB voice is preferred"],["File will not process","Confirm type and size, then try a smaller version"],["Progress disappeared","Check the correct browser/profile and restore a JSON backup"],["Cloud sync failed","Continue locally, verify configuration and select Sync now later"]],[2.0,4.5])

doc.add_heading("Make it part of real life",1)
doc.add_paragraph("Before a meeting, review three relevant words. After reading a report, extract and save only vocabulary you expect to reuse. At the end of the week, download the report and choose one communication habit to improve.")
callout("A practical weekly goal", "Learn five useful words, write two original sentences, complete five reviews and practise one spoken response. Consistency beats volume.")

footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run("Articulate Daily v3.03  •  Customer Guide"); r.font.size=Pt(9); r.font.color.rgb=muted
doc.core_properties.title="Articulate Daily Customer Guide"; doc.core_properties.subject="How to use Articulate Daily effectively"; doc.core_properties.author="Articulate Daily"
doc.save(OUT); print(OUT)
